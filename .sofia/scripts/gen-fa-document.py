#!/usr/bin/env python3
"""
gen-fa-document.py — FA-Agent SOFIA v2.6 (GENERICO INCREMENTAL)
Genera y actualiza el Analisis Funcional acumulativo del proyecto.
- Documento UNICO por proyecto, versionado e incremental sprint a sprint
- 100% dinamico desde fa-index.json — sin datos hardcodeados
- Estructura: Portada → Indice → 1.Ejecutivo → 2.Contexto → 3.Arquitectura
              → 4.Catalogo (por sprint) → 5.RN → 6.Glosario → 7.Cobertura → 8.Historial
- Versionado: doc_version se incrementa en cada ejecucion
- Indice con hipervinculos internos Word clickables (w:hyperlink + w:anchor)
- Compatible con cualquier proyecto SOFIA

Uso:
  python3 .sofia/scripts/gen-fa-document.py

Dependencias:
  pip3 install python-docx --break-system-packages
"""

import os, json, sys
from datetime import datetime, timezone
from pathlib import Path
from collections import defaultdict

# ── Rutas dinamicas (LA-CORE-003) ──────────────────────────────────────────────
_SCRIPT_DIR  = Path(os.path.abspath(__file__)).parent
PROJECT_ROOT = Path(os.path.realpath(_SCRIPT_DIR / '..' / '..'))

def _read_config():
    project, client = 'Proyecto', 'Cliente'
    sofia_ver = '2.6'
    for fpath in [PROJECT_ROOT/'.sofia'/'sofia-config.json', PROJECT_ROOT/'.sofia'/'session.json']:
        if fpath.exists():
            try:
                d = json.loads(fpath.read_text())
                project = d.get('project', project)
                client  = d.get('client',  client)
                sofia_ver = d.get('sofia_version', sofia_ver)
                break
            except Exception: pass
    return project, client, sofia_ver

PROJECT_NAME, CLIENT_NAME, SOFIA_VERSION = _read_config()
DOC_NAME  = f'FA-{PROJECT_NAME}-{CLIENT_NAME}.docx'
FA_DIR    = PROJECT_ROOT / 'docs' / 'functional-analysis'
OUT_PATH  = FA_DIR / DOC_NAME
INDEX_PATH = FA_DIR / 'fa-index.json'
SESSION_PATH = PROJECT_ROOT / '.sofia' / 'session.json'
LOG_PATH  = PROJECT_ROOT / '.sofia' / 'sofia.log'

# ── python-docx ────────────────────────────────────────────────────────────────
try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print('ERROR: python-docx no instalado. Ejecutar:')
    print('  pip3 install python-docx --break-system-packages')
    sys.exit(1)

# ── Colores ─────────────────────────────────────────────────────────────────────
BLUE_DARK  = RGBColor(0x1B, 0x3A, 0x6B)
BLUE_MID   = RGBColor(0x45, 0x69, 0x9E)
ORANGE     = RGBColor(0xC8, 0x4A, 0x14)
GREEN_SOFT = RGBColor(0x5B, 0x88, 0x78)
AMBER      = RGBColor(0xB4, 0x53, 0x09)
GREEN_DARK = RGBColor(0x1A, 0x5C, 0x2A)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
GRAY_TEXT  = RGBColor(0x33, 0x33, 0x33)
GRAY_SUB   = RGBColor(0x66, 0x66, 0x66)
GRAY_LIGHT = RGBColor(0xAA, 0xAA, 0xAA)

# ── Helpers ──────────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color='D9DDE8'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_header_row(table, texts, widths_cm):
    row = table.rows[0]
    for i, (cell, text) in enumerate(zip(row.cells, texts)):
        if widths_cm: cell.width = Cm(widths_cm[i])
        set_cell_bg(cell, '1B3A6B')
        set_cell_borders(cell, '45699E')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9)
        run.font.name = 'Arial'

def add_data_row(table, texts, widths_cm=None, alt=False, highlight=False):
    row = table.add_row()
    bg = 'EFF6FF' if highlight else ('F5F7FA' if alt else 'FFFFFF')
    for i, (cell, text) in enumerate(zip(row.cells, texts)):
        if widths_cm: cell.width = Cm(widths_cm[i])
        set_cell_bg(cell, bg)
        set_cell_borders(cell, 'D9DDE8')
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(9)
        run.font.name = 'Arial'
        run.font.color.rgb = GRAY_TEXT

def heading1(doc, text, bookmark_id=None):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = BLUE_DARK
        run.font.name = 'Arial'
        run.font.size = Pt(14)
    if bookmark_id:
        _add_bookmark(p, bookmark_id)
    return p

def heading2(doc, text, bookmark_id=None):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = BLUE_MID
        run.font.name = 'Arial'
        run.font.size = Pt(12)
    if bookmark_id:
        _add_bookmark(p, bookmark_id)
    return p

def heading3(doc, text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.color.rgb = ORANGE
        run.font.name = 'Arial'
        run.font.size = Pt(11)
    return p

def body(doc, text, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(str(text))
    run.font.size = Pt(10)
    run.font.name = 'Arial'
    run.font.color.rgb = color or GRAY_TEXT
    run.italic = italic
    return p

def bullet_item(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(str(text))
    run.font.size = Pt(9)
    run.font.name = 'Arial'
    run.font.color.rgb = GRAY_TEXT
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    return p

# ── Bookmarks e hipervinculos internos Word ──────────────────────────────────────

# Contador global para IDs de bookmark unicos y secuenciales
_bookmark_counter = [0]

def _next_bid():
    _bookmark_counter[0] += 1
    return str(_bookmark_counter[0])

def _add_bookmark(paragraph, bookmark_id):
    """
    Inserta un bookmark Word en el parrafo para que los hipervinculos del indice
    puedan apuntar a el. Usa IDs secuenciales para evitar colisiones.
    """
    p = paragraph._p
    bid = _next_bid()
    bStart = OxmlElement('w:bookmarkStart')
    bStart.set(qn('w:id'), bid)
    bStart.set(qn('w:name'), bookmark_id)
    bEnd = OxmlElement('w:bookmarkEnd')
    bEnd.set(qn('w:id'), bid)
    p.insert(0, bStart)
    p.append(bEnd)

def add_toc_hyperlink(doc, number, title, anchor, level=1):
    """
    Crea una entrada de indice como hipervinculo interno Word clickable.

    Mecanismo: w:hyperlink con w:anchor apunta al bookmark del heading de destino.
    Al hacer clic en Word o LibreOffice navega directamente a la seccion.
    En Word se activa con Ctrl+Clic (vista normal) o clic directo (vista lectura).
    """
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(level * 0.6 if level > 1 else 0.2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(3 if level > 1 else 5)

    # w:hyperlink con w:anchor = nombre del bookmark de destino
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), anchor)
    hyperlink.set(qn('w:history'), '1')

    # Propiedades del run (estilo hipervinculo)
    run_elem = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Estilo base: rStyle TOC para que Word lo reconozca como entrada de indice
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)

    # Fuente Arial
    rFont = OxmlElement('w:rFonts')
    rFont.set(qn('w:ascii'),  'Arial')
    rFont.set(qn('w:hAnsi'), 'Arial')
    rFont.set(qn('w:cs'),    'Arial')
    rPr.append(rFont)

    # Tamaño
    sz_val = str((10 if level == 1 else 9) * 2)
    rSz = OxmlElement('w:sz')
    rSz.set(qn('w:val'), sz_val)
    rPr.append(rSz)
    rSzCs = OxmlElement('w:szCs')
    rSzCs.set(qn('w:val'), sz_val)
    rPr.append(rSzCs)

    # Color: azul oscuro nivel 1, azul medio nivel 2
    color_val = '1B3A6B' if level == 1 else '45699E'
    rColor = OxmlElement('w:color')
    rColor.set(qn('w:val'), color_val)
    rPr.append(rColor)

    # Subrayado — visual de hipervinculo
    rU = OxmlElement('w:u')
    rU.set(qn('w:val'), 'single')
    rPr.append(rU)

    # Negrita para entradas de nivel 1
    if level == 1:
        rB = OxmlElement('w:b')
        rPr.append(rB)
        rBCs = OxmlElement('w:bCs')
        rPr.append(rBCs)

    run_elem.append(rPr)

    # Texto de la entrada
    label = f'{number.strip()}.  {title}'
    t_elem = OxmlElement('w:t')
    t_elem.text = label
    t_elem.set(qn('xml:space'), 'preserve')
    run_elem.append(t_elem)

    hyperlink.append(run_elem)
    p._p.append(hyperlink)
    return p


# ── Portada ──────────────────────────────────────────────────────────────────────
def build_cover(doc, fa, session, doc_version, now_str):
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ANÁLISIS FUNCIONAL')
    run.bold = True; run.font.size = Pt(28)
    run.font.color.rgb = BLUE_DARK; run.font.name = 'Arial'

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(f'{PROJECT_NAME} — {CLIENT_NAME}')
    run2.bold = True; run2.font.size = Pt(20)
    run2.font.color.rgb = ORANGE; run2.font.name = 'Arial'

    doc.add_paragraph()

    meta = doc.add_table(rows=6, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    pairs = [
        ('Proyecto',         f'{PROJECT_NAME} · {CLIENT_NAME}'),
        ('Versión',          f'v{doc_version}'),
        ('Fecha',            now_str[:10]),
        ('Sprints',          f"S1–S{fa.get('last_sprint',1)} consolidados"),
        ('Funcionalidades',  str(fa.get('total_functionalities', 0))),
        ('Reglas de Negocio',str(fa.get('total_business_rules', 0))),
    ]
    for i, (k, v) in enumerate(pairs):
        lc = meta.rows[i].cells[0]; vc = meta.rows[i].cells[1]
        set_cell_bg(lc, '1B3A6B'); set_cell_bg(vc, 'F5F7FA')
        set_cell_borders(lc, '45699E'); set_cell_borders(vc, 'D9DDE8')
        lc.width = Cm(5); vc.width = Cm(9)
        rl = lc.paragraphs[0].add_run(k)
        rl.bold = True; rl.font.color.rgb = WHITE
        rl.font.size = Pt(10); rl.font.name = 'Arial'
        rv = vc.paragraphs[0].add_run(v)
        rv.font.color.rgb = BLUE_DARK; rv.font.size = Pt(10)
        rv.font.name = 'Arial'; rv.bold = True

    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(
        f'Generado por FA-Agent · SOFIA v{SOFIA_VERSION} · '
        f'{now_str[:16].replace("T"," ")} UTC'
    )
    run3.font.size = Pt(8); run3.font.color.rgb = GRAY_LIGHT; run3.font.name = 'Arial'

    doc.add_page_break()


# ── Índice con hipervínculos internos clickables ──────────────────────────────────
def build_toc(doc, fa):
    """
    Genera la tabla de contenidos con hipervinculos internos Word.

    Cada entrada usa w:hyperlink + w:anchor apuntando al bookmark_id del heading
    correspondiente. En Word: Ctrl+Clic (modo edicion) o clic directo (modo lectura).
    En LibreOffice Writer: clic directo siempre.

    Los anchor names DEBEN coincidir exactamente con los bookmark_id declarados
    en los headings de cada seccion (heading1/heading2 con bookmark_id=...).
    """
    heading1(doc, 'ÍNDICE DE CONTENIDOS', bookmark_id='toc')
    doc.add_paragraph()

    # Calcular sprints para las entradas 4.N del catalogo
    sprints_seen = sorted(set(
        str(f.get('sprint', f.get('feat', '?')))
        for f in fa.get('functionalities', [])
    ), key=lambda x: (int(x.split('-')[0]) if x.split('-')[0].isdigit() else 999))

    feat_by_sprint = defaultdict(list)
    for f in fa.get('functionalities', []):
        sp  = str(f.get('sprint', '?'))
        fea = f.get('feature', f.get('feat', '?'))
        feat_by_sprint[sp].append(fea)

    # Entradas del indice: (numero, titulo, anchor_bookmark, nivel)
    # CRITICO: anchor debe coincidir EXACTAMENTE con bookmark_id del heading
    toc_entries = [
        ('1',   'RESUMEN EJECUTIVO',           'sec1',   1),
        ('2',   'CONTEXTO DE NEGOCIO',         'sec2',   1),
        ('2.1', 'Descripción del proyecto',    'sec2_1', 2),
        ('2.2', 'Actores del sistema',         'sec2_2', 2),
        ('2.3', 'Marco tecnológico',           'sec2_3', 2),
        ('3',   'ARQUITECTURA FUNCIONAL',      'sec3',   1),
        ('3.1', 'Módulos del sistema',         'sec3_1', 2),
        ('3.2', 'Mapa de dependencias',        'sec3_2', 2),
        ('4',   'CATÁLOGO DE FUNCIONALIDADES', 'sec4',   1),
    ]

    for i, sp in enumerate(sprints_seen, 1):
        feats = sorted(set(feat_by_sprint[sp]))
        feat_label = ' · '.join(feats) if feats else '—'
        anchor = f'sprint_{sp}'   # coincide con bookmark_id en build_section4
        toc_entries.append((f'4.{i}', f'Sprint {sp} — {feat_label}', anchor, 2))

    toc_entries += [
        ('5', 'REGLAS DE NEGOCIO CONSOLIDADAS',  'sec5', 1),
        ('6', 'GLOSARIO DEL DOMINIO',            'sec6', 1),
        ('7', 'MATRIZ DE COBERTURA FUNCIONAL',   'sec7', 1),
        ('8', 'HISTORIAL DE CAMBIOS',            'sec8', 1),
    ]

    # Separador superior
    p_sep = doc.add_paragraph()
    r_sep = p_sep.add_run('─' * 74)
    r_sep.font.size = Pt(8); r_sep.font.color.rgb = BLUE_MID; r_sep.font.name = 'Arial'

    # Generar hipervinculos
    for num, title, anchor, level in toc_entries:
        add_toc_hyperlink(doc, num, title, anchor, level)

    # Separador inferior
    p_sep2 = doc.add_paragraph()
    r_sep2 = p_sep2.add_run('─' * 74)
    r_sep2.font.size = Pt(8); r_sep2.font.color.rgb = BLUE_MID; r_sep2.font.name = 'Arial'

    # Nota de uso
    p_note = doc.add_paragraph()
    r_note = p_note.add_run(
        'ℹ  Haga clic en cualquier entrada del índice para navegar directamente a la sección. '
        '(En Word: Ctrl+Clic en modo edición · Clic directo en modo lectura)'
    )
    r_note.font.size = Pt(8); r_note.font.color.rgb = GRAY_SUB
    r_note.font.name = 'Arial'; r_note.italic = True

    doc.add_page_break()


# ── Sección 1: Resumen Ejecutivo ─────────────────────────────────────────────────
def build_section1(doc, fa, session):
    heading1(doc, '1. RESUMEN EJECUTIVO', bookmark_id='sec1')

    total_f  = fa.get('total_functionalities', 0)
    total_rn = fa.get('total_business_rules', 0)
    last_sp  = fa.get('last_sprint', 1)
    last_feat = fa.get('last_feat', '—')

    body(doc,
        f'Este documento recoge el Análisis Funcional acumulativo del proyecto {PROJECT_NAME} '
        f'para {CLIENT_NAME}, gestionado bajo la metodología SOFIA v{SOFIA_VERSION} con CMMI '
        f'Nivel 3 y Scrumban. El documento es un artefacto vivo que se actualiza '
        f'incrementalmente en cada sprint.')

    doc.add_paragraph()
    tw = [3.5, 3.5, 3.5, 4.2]
    t = doc.add_table(rows=1, cols=4)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    add_header_row(t, ['Funcionalidades', 'Reglas de Negocio', 'Sprints consolidados', 'Última feature'], tw)
    add_data_row(t, [str(total_f), str(total_rn), f'S1–S{last_sp}', last_feat], tw)

    doc.add_paragraph()
    heading2(doc, '1.1 Estado por sprint')

    sh = session.get('sprint_history', {})
    sw = [1.5, 3.5, 1.5, 2.0, 2.0, 2.0, 1.8]
    t2 = doc.add_table(rows=1, cols=7)
    t2.alignment = WD_TABLE_ALIGNMENT.LEFT
    add_header_row(t2, ['Sprint', 'Feature', 'SP', 'Tests', 'Cobertura', 'Release', 'Estado'], sw)
    for i, (k, h) in enumerate(sorted(
        sh.items(),
        key=lambda x: int(x[0].replace('sprint_','').split('_')[0]) if x[0].replace('sprint_','').split('_')[0].isdigit() else 999
    )):
        sp_num = k.replace('sprint_', 'S')
        add_data_row(t2, [
            sp_num,
            h.get('feature', '—'),
            str(h.get('sp', 0)),
            str(h.get('tests', '—')),
            f"{h.get('cov','—')}%" if h.get('cov') else '—',
            h.get('rel', '—'),
            h.get('status', '—').upper(),
        ], sw, alt=(i % 2 == 1))

    doc.add_paragraph()
    heading2(doc, '1.2 Leyenda de evidencias')
    body(doc, 'Los artefactos de este documento se clasifican por nivel de evidencia:')
    bullet_item(doc, 'SPRINT-N: Documentado con evidencia directa del sprint cerrado — artefactos verificados.')
    bullet_item(doc, 'RECONSTRUIDO: Inferido desde código fuente, migraciones o retrospectiva — evidencia indirecta.')
    bullet_item(doc, 'PLANNED: Funcionalidad planificada en el sprint activo — pendiente de entrega.')

    doc.add_page_break()


# ── Sección 2: Contexto de Negocio ──────────────────────────────────────────────
def build_section2(doc, fa, session):
    heading1(doc, '2. CONTEXTO DE NEGOCIO', bookmark_id='sec2')

    heading2(doc, '2.1 Descripción del proyecto', bookmark_id='sec2_1')
    desc = fa.get('description',
        f'{PROJECT_NAME} es un sistema desarrollado para {CLIENT_NAME} bajo la metodología '
        f'SOFIA con CMMI Nivel 3. El desarrollo sigue ciclos Scrumban de 2 semanas con '
        f'pipeline de calidad automatizado (17 steps, 9 gates HITL).')
    body(doc, desc)

    doc.add_paragraph()
    heading2(doc, '2.2 Actores del sistema', bookmark_id='sec2_2')
    actors = fa.get('actors', [
        {'rol': 'Usuario final',  'desc': 'Usuario principal del sistema', 'acceso': 'Web / API'},
        {'rol': 'Administrador',  'desc': 'Gestión de configuración y catálogos', 'acceso': 'Web / API'},
        {'rol': 'Sistema',        'desc': 'Procesos automáticos y schedulers', 'acceso': 'Interno'},
    ])
    aw = [3.5, 6.5, 4.7]
    t = doc.add_table(rows=1, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    add_header_row(t, ['Rol', 'Descripción', 'Acceso'], aw)
    for i, a in enumerate(actors):
        add_data_row(t, [
            a.get('rol', a.get('name', '—')),
            a.get('desc', a.get('description', '—')),
            a.get('acceso', a.get('access', 'API REST / Web')),
        ], aw, alt=(i % 2 == 1))

    doc.add_paragraph()
    heading2(doc, '2.3 Marco tecnológico', bookmark_id='sec2_3')
    cfg_path = PROJECT_ROOT / '.sofia' / 'sofia-config.json'
    stack_be = stack_fe = stack_db = stack_inf = '—'
    if cfg_path.exists():
        try:
            cfg = json.loads(cfg_path.read_text())
            stk = cfg.get('stack', {})
            stack_be  = ' · '.join(stk.get('backend',  [])) or '—'
            stack_fe  = ' · '.join(stk.get('frontend', [])) or '—'
            stack_db  = ' · '.join(stk.get('database', [])) or '—'
            stack_inf = ' · '.join(stk.get('infra',    [])) or 'Docker · Docker Compose'
        except Exception: pass

    rw = [4.0, 10.7]
    t2 = doc.add_table(rows=1, cols=2)
    t2.alignment = WD_TABLE_ALIGNMENT.LEFT
    add_header_row(t2, ['Capa', 'Tecnología'], rw)
    for i, (capa, tec) in enumerate([
        ('Backend',       stack_be),
        ('Frontend',      stack_fe),
        ('Base de datos', stack_db),
        ('Infraestructura', stack_inf),
        ('Metodología',   f'SOFIA v{SOFIA_VERSION} · Scrumban · CMMI L3 · 17 steps pipeline'),
    ]):
        add_data_row(t2, [capa, tec], rw, alt=(i % 2 == 1))

    regs = fa.get('regulations', [])
    if regs:
        doc.add_paragraph()
        heading2(doc, '2.4 Marco regulatorio', bookmark_id='sec2_4')
        rw2 = [3.0, 11.7]
        t3 = doc.add_table(rows=1, cols=2)
        t3.alignment = WD_TABLE_ALIGNMENT.LEFT
        add_header_row(t3, ['Regulación/Estándar', 'Ámbito'], rw2)
        for i, reg in enumerate(regs):
            if isinstance(reg, dict):
                add_data_row(t3, [reg.get('id','—'), reg.get('desc','—')], rw2, alt=(i%2==1))
            else:
                add_data_row(t3, [str(reg), 'Ver funcionalidades afectadas'], rw2, alt=(i%2==1))

    doc.add_page_break()


# ── Sección 3: Arquitectura Funcional ───────────────────────────────────────────
def build_section3(doc, fa):
    heading1(doc, '3. ARQUITECTURA FUNCIONAL', bookmark_id='sec3')
    heading2(doc, '3.1 Módulos del sistema', bookmark_id='sec3_1')

    module_map = defaultdict(lambda: {'features': set(), 'sprints': set(), 'funcs': 0})
    for f in fa.get('functionalities', []):
        mod = f.get('module', f.get('area', 'general'))
        module_map[mod]['features'].add(f.get('feature', f.get('feat', '?')))
        module_map[mod]['sprints'].add(str(f.get('sprint', '?')))
        module_map[mod]['funcs'] += 1

    mw = [3.0, 4.5, 2.5, 2.0, 2.7]
    t = doc.add_table(rows=1, cols=5)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    add_header_row(t, ['Módulo', 'Features', 'Sprints', 'FA count', 'Estado'], mw)
    for i, (mod, info) in enumerate(sorted(module_map.items())):
        feats   = ', '.join(sorted(info['features']))
        sprints = ', '.join(sorted(info['sprints'], key=lambda x: int(x) if x.isdigit() else 999))
        add_data_row(t, [mod, feats[:40], sprints, str(info['funcs']), 'ENTREGADO'], mw, alt=(i%2==1))

    doc.add_paragraph()
    heading2(doc, '3.2 Mapa de dependencias', bookmark_id='sec3_2')
    body(doc,
        'Las dependencias entre módulos se documentan en los documentos de arquitectura '
        '(HLD/LLD) de cada sprint. El flujo general sigue Clean Architecture / Onion con '
        'separación estricta entre Domain, Application e Infrastructure.')

    doc.add_page_break()


# ── Sección 4: Catálogo de Funcionalidades ───────────────────────────────────────
def build_section4(doc, fa):
    heading1(doc, '4. CATÁLOGO DE FUNCIONALIDADES', bookmark_id='sec4')
    body(doc,
        f'Catálogo acumulativo de {fa.get("total_functionalities",0)} funcionalidades '
        f'distribuidas en {fa.get("last_sprint",1)} sprint(s). '
        f'Se documenta el comportamiento funcional en lenguaje de negocio.')

    by_sprint = defaultdict(list)
    for f in fa.get('functionalities', []):
        sp = str(f.get('sprint', '?'))
        by_sprint[sp].append(f)

    def sprint_sort_key(s):
        try: return int(s.split('-')[0])
        except: return 999

    sprints_sorted = sorted(by_sprint.keys(), key=sprint_sort_key)
    rn_map = {r['id']: r for r in fa.get('business_rules', [])}

    for idx, sp in enumerate(sprints_sorted, 1):
        funcs = by_sprint[sp]
        feats = sorted(set(f.get('feature', f.get('feat', '?')) for f in funcs))
        feat_label = ' · '.join(feats)

        # bookmark_id='sprint_{sp}' coincide con el anchor del TOC
        heading2(doc, f'4.{idx}  Sprint {sp} — {feat_label}', bookmark_id=f'sprint_{sp}')

        src_labels = set(f.get('source', f.get('src', 'SPRINT')) for f in funcs)
        src = ', '.join(src_labels)
        src_color = (GREEN_SOFT if 'SPRINT' in src or 'EVIDENCIA' in src
                     else AMBER if 'RECON' in src.upper() else BLUE_MID)
        p = doc.add_paragraph()
        r = p.add_run(f'Evidencia: {src}')
        r.font.size = Pt(8); r.font.name = 'Arial'
        r.font.color.rgb = src_color; r.italic = True

        for f in funcs:
            fid    = f.get('id', '—')
            fname  = f.get('name', f.get('title', '—'))
            fdesc  = f.get('description', '—')
            frns   = f.get('business_rules', [])
            status = f.get('status', 'DELIVERED')

            heading3(doc, f'{fid} — {fname}')

            p_st = doc.add_paragraph()
            r_st = p_st.add_run(f'Estado: {status}')
            r_st.font.size = Pt(9); r_st.font.name = 'Arial'; r_st.bold = True
            r_st.font.color.rgb = (GREEN_DARK if 'DELIV' in status.upper()
                                   else AMBER if 'PLAN' in status.upper() else BLUE_MID)

            body(doc, fdesc)

            if frns:
                p_rn = doc.add_paragraph()
                r_rn = p_rn.add_run('Reglas de negocio aplicables:')
                r_rn.bold = True; r_rn.font.size = Pt(10)
                r_rn.font.name = 'Arial'; r_rn.font.color.rgb = BLUE_MID
                for rn_id in frns:
                    rn = rn_map.get(rn_id, {})
                    rn_desc = rn.get('desc', rn.get('description', rn_id))
                    bullet_item(doc, f'[{rn_id}] {rn_desc}')

        if idx < len(sprints_sorted):
            doc.add_page_break()

    doc.add_page_break()


# ── Sección 5: Reglas de Negocio Consolidadas ────────────────────────────────────
def build_section5(doc, fa):
    heading1(doc, '5. REGLAS DE NEGOCIO CONSOLIDADAS', bookmark_id='sec5')
    body(doc, f'Catálogo consolidado de {fa.get("total_business_rules",0)} reglas de negocio agrupadas por módulo.')

    by_module = defaultdict(list)
    for rn in fa.get('business_rules', []):
        mod = rn.get('module', rn.get('area', rn.get('fa', 'general')))
        by_module[mod].append(rn)

    for mod, rules in sorted(by_module.items()):
        heading2(doc, f'  {mod.upper()}')
        rnw = [2.0, 12.7]
        t = doc.add_table(rows=1, cols=2)
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        add_header_row(t, ['ID', 'Regla de Negocio'], rnw)
        for i, rn in enumerate(rules):
            add_data_row(t, [rn.get('id','—'), rn.get('desc', rn.get('description','—'))], rnw, alt=(i%2==1))
        doc.add_paragraph()

    doc.add_page_break()


# ── Sección 6: Glosario ──────────────────────────────────────────────────────────
def build_section6(doc, fa):
    heading1(doc, '6. GLOSARIO DEL DOMINIO', bookmark_id='sec6')
    body(doc, 'Glosario acumulativo de términos de negocio y técnicos. Se amplía cada sprint.')

    glossary = fa.get('glossary', [])
    if not glossary:
        glossary = [
            {'term': 'SOFIA',     'def': 'Software Factory IA de Experis — orquesta agentes especializados via MCP'},
            {'term': 'CMMI L3',   'def': 'Capability Maturity Model Integration Nivel 3 — procesos definidos y gestionados'},
            {'term': 'Scrumban',  'def': 'Metodología híbrida Scrum + Kanban — sprints de 2 semanas con WIP limits'},
            {'term': 'Gate HITL', 'def': 'Human-In-The-Loop gate — punto de aprobación explícita por rol definido'},
            {'term': 'FA',        'def': 'Functional Analysis — análisis funcional en lenguaje de negocio'},
            {'term': 'RN',        'def': 'Regla de Negocio — restricción o comportamiento obligatorio del sistema'},
            {'term': 'US',        'def': 'User Story — historia de usuario en formato Agile'},
            {'term': 'SP',        'def': 'Story Points — unidad relativa de estimación de esfuerzo'},
            {'term': 'API REST',  'def': 'Interfaz basada en HTTP con recursos y verbos estándar'},
            {'term': 'JWT',       'def': 'JSON Web Token — token de autenticación sin estado, firmado criptográficamente'},
        ]

    gw = [3.5, 11.2]
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    add_header_row(t, ['Término', 'Definición'], gw)
    for i, entry in enumerate(glossary):
        if isinstance(entry, dict):
            term = entry.get('term', entry.get('name', '—'))
            defn = entry.get('def', entry.get('definition', entry.get('desc', '—')))
        else:
            term, defn = str(entry), '—'
        add_data_row(t, [term, defn], gw, alt=(i%2==1))

    doc.add_page_break()


# ── Sección 7: Matriz de Cobertura ───────────────────────────────────────────────
def build_section7(doc, fa):
    heading1(doc, '7. MATRIZ DE COBERTURA FUNCIONAL', bookmark_id='sec7')
    body(doc, f'Trazabilidad completa de las {fa.get("total_functionalities",0)} funcionalidades.')

    mxw = [2.0, 2.5, 4.5, 1.5, 2.0, 2.0, 2.2]
    t = doc.add_table(rows=1, cols=7)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    add_header_row(t, ['ID', 'Feature', 'Funcionalidad', 'Sprint', 'Versión', 'Estado', 'Evidencia'], mxw)

    for i, f in enumerate(fa.get('functionalities', [])):
        fid    = f.get('id', '—')
        feat   = f.get('feature', f.get('feat', '—'))
        fname  = f.get('name', f.get('title', '—'))
        sp     = str(f.get('sprint', '?'))
        ver    = f.get('version', '—')
        status = f.get('status', 'DELIVERED')
        src    = f.get('source', f.get('src', 'SPRINT'))
        hi     = 'DIRECTA' in src.upper() or 'SPRINT' in src.upper()
        add_data_row(t, [fid, feat, fname[:50], f'S{sp}', ver, status, src[:20]],
                     mxw, alt=(i%2==1), highlight=hi)

    doc.add_page_break()


# ── Sección 8: Historial de cambios ─────────────────────────────────────────────
def build_section8(doc, fa, doc_version, now_str):
    heading1(doc, '8. HISTORIAL DE CAMBIOS DEL DOCUMENTO', bookmark_id='sec8')
    body(doc, 'Registro de versiones. El documento se actualiza incrementalmente en cada sprint.')

    hw = [1.5, 2.0, 3.5, 8.0]
    t = doc.add_table(rows=1, cols=4)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    add_header_row(t, ['Versión', 'Fecha', 'Sprint / Feature', 'Descripción del cambio'], hw)

    history = fa.get('doc_history', [])
    if not history:
        history = [{'version': '1.0', 'date': fa.get('generated_at', now_str)[:10],
                    'sprint': f"S1–S{fa.get('last_sprint',1)}",
                    'desc': f'Creación inicial. {fa.get("total_functionalities",0)} func, '
                            f'{fa.get("total_business_rules",0)} RN.'}]

    for i, h in enumerate(history):
        add_data_row(t, [
            h.get('version','—'), h.get('date','—')[:10],
            h.get('sprint','—'), h.get('desc', h.get('description','—')),
        ], hw, alt=(i%2==1))

    last_ver = history[-1].get('version','0') if history else '0'
    if last_ver != doc_version:
        add_data_row(t, [
            doc_version, now_str[:10],
            f"S1–S{fa.get('last_sprint',1)} / {fa.get('last_feat','—')}",
            f'Actualización automática. {fa.get("total_functionalities",0)} func, '
            f'{fa.get("total_business_rules",0)} RN. SOFIA v{SOFIA_VERSION}.',
        ], hw, alt=(len(history)%2==1))

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run(
        f'Documento generado por FA-Agent — SOFIA Software Factory v{SOFIA_VERSION} · '
        f'{PROJECT_NAME} · {CLIENT_NAME}'
    )
    r.font.size = Pt(8); r.font.name = 'Arial'; r.font.color.rgb = GRAY_LIGHT


# ── Header / Footer ─────────────────────────────────────────────────────────────
def set_header_footer(doc, doc_version, now_str):
    for section in doc.sections:
        section.page_width    = Cm(21)
        section.page_height   = Cm(29.7)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)

        hdr = section.header
        p = hdr.paragraphs[0]; p.clear()
        run = p.add_run(f'Análisis Funcional — {PROJECT_NAME} · {CLIENT_NAME} · v{doc_version}')
        run.font.size = Pt(8); run.font.name = 'Arial'; run.font.color.rgb = BLUE_MID
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        ftr = section.footer
        pf = ftr.paragraphs[0]; pf.clear()
        run_f = pf.add_run(f'SOFIA v{SOFIA_VERSION} · Confidencial · Experis · {now_str[:10]}')
        run_f.font.size = Pt(8); run_f.font.name = 'Arial'; run_f.font.color.rgb = GRAY_LIGHT


# ── Main ─────────────────────────────────────────────────────────────────────────
def generate(out_path):
    if not INDEX_PATH.exists():
        print(f'ERROR: fa-index.json no encontrado en {INDEX_PATH}')
        sys.exit(1)
    fa = json.loads(INDEX_PATH.read_text())

    session = {}
    if SESSION_PATH.exists():
        try: session = json.loads(SESSION_PATH.read_text())
        except Exception: pass

    now     = datetime.now(timezone.utc)
    now_str = now.isoformat()

    old_ver = fa.get('doc_version', '0.0')
    try:
        major, minor = map(int, str(old_ver).split('.'))
        doc_version = f'{major}.{minor + 1}'
    except Exception:
        doc_version = '1.0'

    print(f'Generando {DOC_NAME} v{doc_version}...')
    print(f'  Funcionalidades: {fa.get("total_functionalities",0)}')
    print(f'  Reglas de negocio: {fa.get("total_business_rules",0)}')
    print(f'  Sprints: S1–S{fa.get("last_sprint",1)}')
    print(f'  Indice: hipervinculos internos Word clickables')

    # Resetear contador de bookmarks en cada ejecucion
    _bookmark_counter[0] = 0

    doc = Document()
    set_header_footer(doc, doc_version, now_str)

    build_cover(doc, fa, session, doc_version, now_str)
    build_toc(doc, fa)
    build_section1(doc, fa, session)
    build_section2(doc, fa, session)
    build_section3(doc, fa)
    build_section4(doc, fa)
    build_section5(doc, fa)
    build_section6(doc, fa)
    build_section7(doc, fa)
    build_section8(doc, fa, doc_version, now_str)

    FA_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    size_kb = round(os.path.getsize(str(out_path)) / 1024, 1)
    print(f'OK Guardado: {out_path} ({size_kb} KB)')

    # Actualizar fa-index.json
    fa['doc_version'] = doc_version
    fa['doc_path']    = f'docs/functional-analysis/{DOC_NAME}'
    fa['document']    = DOC_NAME
    fa['updated_at']  = now_str
    if 'doc_history' not in fa:
        fa['doc_history'] = []
    fa['doc_history'].append({
        'version': doc_version,
        'date':    now_str[:10],
        'sprint':  f"S1–S{fa.get('last_sprint',1)} / {fa.get('last_feat','—')}",
        'desc':    (f'Actualización automática gate-8b. '
                    f'{fa.get("total_functionalities",0)} func, '
                    f'{fa.get("total_business_rules",0)} RN. '
                    f'Indice clickable v2.6. SOFIA v{SOFIA_VERSION}.'),
    })
    INDEX_PATH.write_text(json.dumps(fa, indent=2, ensure_ascii=False))
    print(f'OK fa-index.json actualizado — doc_version: {doc_version}')

    # Actualizar session.json
    if SESSION_PATH.exists():
        try:
            s = json.loads(SESSION_PATH.read_text())
            if 'fa_agent' not in s: s['fa_agent'] = {}
            s['fa_agent'].update({
                'last_gate': '8b', 'last_updated': now_str,
                'doc_version': doc_version,
                'doc_path': f'docs/functional-analysis/{DOC_NAME}',
                'functionalities': fa.get('total_functionalities', 0),
                'business_rules': fa.get('total_business_rules', 0),
                'generator': 'python-docx',
                'docx_verified': True, 'docx_size_kb': size_kb,
                'last_sprint_consolidated': fa.get('last_sprint', 1),
            })
            s['updated_at'] = now_str
            SESSION_PATH.write_text(json.dumps(s, indent=2, ensure_ascii=False))
            print('OK session.json actualizado')
        except Exception as e:
            print(f'WARN: session.json no actualizado: {e}')

    # Registrar en sofia.log
    if LOG_PATH.exists():
        try:
            with open(str(LOG_PATH), 'a') as lf:
                lf.write(
                    f'[{now_str}] [FA-AGENT] [GATE-8b] GENERATED → {DOC_NAME} '
                    f'v{doc_version} | {size_kb}KB | {fa.get("total_functionalities",0)} func '
                    f'| {fa.get("total_business_rules",0)} RN | TOC-hyperlinks | '
                    f'python-docx | SOFIA v{SOFIA_VERSION}\n'
                )
        except Exception: pass

    return out_path, size_kb, doc_version


if __name__ == '__main__':
    out, size, ver = generate(OUT_PATH)
    print(f'\nOK {DOC_NAME} v{ver} — {size}KB → {out}')
